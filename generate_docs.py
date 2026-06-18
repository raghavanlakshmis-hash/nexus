from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────────
def heading1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x53, 0x4A, 0xB7)

def heading2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x3C, 0x34, 0x89)

def heading3(text):
    doc.add_heading(text, level=3)

def body(text):
    doc.add_paragraph(text)

def bullet(text):
    doc.add_paragraph(text, style="List Bullet")

def code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.5)

def table_2col(rows, header=None):
    t = doc.add_table(rows=len(rows) + (1 if header else 0), cols=2)
    t.style = "Table Grid"
    if header:
        hdr = t.rows[0].cells
        hdr[0].text = header[0]
        hdr[1].text = header[1]
        for cell in hdr:
            for run in cell.paragraphs[0].runs:
                run.font.bold = True
    for i, (a, b) in enumerate(rows):
        row = t.rows[i + (1 if header else 0)].cells
        row[0].text = a
        row[1].text = b
    doc.add_paragraph()

# ── TITLE PAGE ───────────────────────────────────────────────────────────────
title = doc.add_heading("Nexus", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x53, 0x4A, 0xB7)

subtitle = doc.add_paragraph("Post-Hospital Recovery Co-Pilot")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)

meta = doc.add_paragraph("Week 3 — Agentic AI Systems | MasteringGenAI")
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("Author: Lakshmi Raghavan")
doc.add_paragraph("Date: June 2026")
doc.add_page_break()

# ── 1. PROJECT OVERVIEW ──────────────────────────────────────────────────────
heading1("1. Project Overview")
body(
    "Nexus is a multi-agent AI system that transforms a hospital discharge summary PDF "
    "into an active 30-day recovery co-pilot. One in five patients is readmitted within "
    "30 days of discharge — often because they leave with a paper instruction sheet they "
    "don't fully understand and no system monitoring their recovery."
)
body(
    "Nexus closes that gap. It reads the discharge PDF, extracts structured clinical data, "
    "checks for medication interactions, generates a plain-language care plan, and then "
    "monitors the patient daily through voice or typed check-ins — escalating to humans "
    "automatically when something goes wrong."
)

heading2("Core Value Proposition")
bullet("Discharge PDF → structured care plan in under 2 minutes")
bullet("Daily voice check-ins with GREEN / YELLOW / RED classification")
bullet("Tier-based escalation: draft message → auto-notify emergency contact → show 911")
bullet("Nothing leaves the system without human approval (human-in-the-loop)")
bullet("30-day persistent memory via Pinecone vector store")

doc.add_page_break()

# ── 2. ARCHITECTURE ──────────────────────────────────────────────────────────
heading1("2. System Architecture")
body(
    "Nexus is built as a LangGraph state machine with five specialized agents coordinated "
    "by a central orchestrator. All agents share a single typed state object (RecoveryState) "
    "that persists across the entire 30-day recovery period."
)

heading2("Agent Graph")
body("Two separate LangGraph graphs handle the two phases of the system:")
bullet("Intake Graph: Intake Agent → Care Plan Agent (runs once on Day 0)")
bullet("Monitoring Graph: Monitoring Agent → Escalation Agent (conditional) → Admin Agent (runs daily)")

heading2("Agent Descriptions")

heading3("Intake Agent")
body(
    "Parses the discharge summary PDF using PyMuPDF, sends the raw text to Claude "
    "claude-sonnet-4-6 with a structured extraction prompt, and populates the RecoveryState "
    "with diagnosis, medications, appointments, warning signs, and dietary/activity "
    "restrictions. Also auto-populates the hospitalization history record and stores "
    "the discharge summary in Pinecone for later retrieval."
)

heading3("Care Plan Agent")
body(
    "Receives the structured data from the Intake Agent. First checks all medications "
    "against the OpenFDA drug label API for interactions and boxed warnings. Flagged "
    "medications are routed to the human approval queue and excluded from the care plan. "
    "Then calls Claude to generate a plain-language, patient-facing care plan with a "
    "day-by-day checklist, medication schedule, and warning signs. Also calls Nebius "
    "Token Factory (Meta-Llama-3.1-70B-Instruct-fast) in parallel to satisfy the rubric "
    "requirement — the Nebius result is logged but the Claude output is used for the plan."
)

heading3("Monitoring Agent")
body(
    "Runs daily. Generates personalized check-in questions based on the patient's diagnosis "
    "(e.g. CHF patients get daily weight and swelling questions). Accepts responses via "
    "ElevenLabs voice-to-text or a typed form. Sends responses to Claude for "
    "GREEN / YELLOW / RED classification using strict rules: any ER warning sign = RED, "
    "three consecutive YELLOW = escalate to RED. Stores each check-in in Pinecone."
)

heading3("Escalation Agent")
body(
    "Handles RED classifications with a three-tier response system: "
    "TIER_1 (urgent, non-life-threatening) drafts a provider message for human approval. "
    "TIER_2 (potentially life-threatening) auto-notifies the emergency contact via SMS "
    "(Twilio) and generates an ER handoff summary. "
    "TIER_3 (actively life-threatening) displays a 911 screen immediately and auto-notifies "
    "the emergency contact — no human approval step."
)

heading3("Admin Agent")
body(
    "Handles non-clinical tasks: 48-hour appointment reminders, weekly recovery summaries "
    "emailed to the caregiver, and daily task tracking. Runs at the end of every monitoring "
    "cycle."
)

heading3("Orchestrator")
body(
    "Builds and compiles the two LangGraph graphs. Uses conditional edges to route between "
    "agents based on state: if intake PDF fails → END, if monitoring = RED → escalation_agent, "
    "else → admin_agent."
)

doc.add_page_break()

# ── 3. TECH STACK ────────────────────────────────────────────────────────────
heading1("3. Tech Stack")

table_2col([
    ("LangGraph", "Multi-agent state machine orchestration"),
    ("Claude claude-sonnet-4-6 (Anthropic)", "Document extraction, care plan generation, check-in classification, transcript parsing"),
    ("Nebius Token Factory", "Parallel care plan generation (Meta-Llama-3.1-70B-Instruct-fast) — rubric requirement"),
    ("Pinecone", "Persistent vector memory across 30-day recovery period (1024-dim, cosine)"),
    ("ElevenLabs STT (Scribe v1)", "Voice check-in transcription"),
    ("OpenFDA API", "Medication interaction and boxed warning checks"),
    ("PyMuPDF (fitz)", "PDF text extraction from discharge summaries"),
    ("Streamlit", "Patient-facing web UI (7 pages)"),
    ("Twilio", "SMS escalation notifications to emergency contacts"),
    ("Gmail SMTP", "Email notifications and weekly summaries"),
    ("n8n", "Daily 8am check-in reminder scheduler"),
    ("fpdf2", "Synthetic test PDF generation"),
    ("python-dotenv", "Environment variable management"),
], header=("Component", "Role"))

doc.add_page_break()

# ── 4. DESIGN DECISIONS ──────────────────────────────────────────────────────
heading1("4. Key Design Decisions")

heading2("4.1 Single Shared State Object")
body(
    "All agents read from and write to a single RecoveryState TypedDict. This was chosen "
    "over agent-specific state objects because recovery data is deeply interconnected — "
    "the care plan depends on intake data, monitoring depends on the care plan, and "
    "escalation depends on monitoring. A shared state eliminates serialization overhead "
    "and makes the system easier to debug."
)

heading2("4.2 Two Separate Graphs (Not One)")
body(
    "Intake + care plan generation happens once. Monitoring happens daily for 30 days. "
    "Combining them into one graph would require complex looping logic. Two separate "
    "graphs with a shared state object is cleaner — the Streamlit UI invokes the right "
    "graph at the right time."
)

heading2("4.3 Hash-based Pinecone Embeddings")
body(
    "Real embedding models (e.g. text-embedding-3-small) cost money per call and add "
    "latency. For this demo, a deterministic MD5 hash-based 1024-dim sparse vector is "
    "used. This means similarity search is approximate but the store/retrieve pipeline "
    "is fully functional. A production system would swap in a real embedding model with "
    "one line change in tools/pinecone_store.py."
)
body("Note: Pinecone index configured with 1024 dimensions (not 1536) as that was the available option.")

heading2("4.4 Human-in-the-Loop for All Outbound Communications")
body(
    "No message to a provider, pharmacist, or external party is sent without the patient "
    "or caregiver reviewing and approving it first — except TIER_2 and TIER_3 escalations "
    "where time criticality overrides the approval step (and consent was obtained upfront). "
    "This was a deliberate safety decision given the medical context."
)

heading2("4.5 Nebius as Parallel Path (Not Fallback)")
body(
    "The rubric requires at least one Nebius Token Factory call. Rather than using it as "
    "a fallback (which would mean it only runs on Claude failure), it runs in parallel "
    "during care plan generation. Claude's output is used for the actual plan; Nebius "
    "output is logged. This satisfies the rubric while keeping the patient experience "
    "consistent."
)

heading2("4.6 Voice-First with Typed Fallback")
body(
    "ElevenLabs STT (Scribe v1) is the primary check-in method. If the transcription "
    "fails or the patient prefers typing, a full typed form is available as a fallback. "
    "Claude parses the free-form voice transcript into structured responses matching "
    "the question schema — patients don't need to answer in a specific format."
)

heading2("4.7 Medication Questions as Individual Checkboxes")
body(
    "The original design used a single yes/no question for all medications. This was "
    "changed to individual checkboxes per medication after testing showed it was impossible "
    "to track which specific medication was missed. Individual checkboxes also improve "
    "data quality for the provider summary."
)

doc.add_page_break()

# ── 5. DATA FLOW ─────────────────────────────────────────────────────────────
heading1("5. Data Flow")

heading2("Day 0 — Onboarding")
bullet("Patient uploads discharge PDF and fills in personal details")
bullet("Intake Agent: PDF → PyMuPDF → raw text → Claude extraction → RecoveryState")
bullet("Care Plan Agent: OpenFDA check → Claude care plan → Nebius parallel call → RecoveryState")
bullet("Discharge summary chunked and stored in Pinecone (namespace: patient_{id})")
bullet("Care plan displayed across 4 tabs in Streamlit")

heading2("Day 1–30 — Daily Check-in")
bullet("Patient speaks or types their check-in responses")
bullet("ElevenLabs STT transcribes voice → Claude parses transcript → structured responses")
bullet("Monitoring Agent: responses + history → Claude classification → GREEN/YELLOW/RED")
bullet("Check-in record stored in Pinecone")
bullet("RED → Escalation Agent (tier determination) → appropriate response")
bullet("Admin Agent: appointment reminders, weekly summaries")

doc.add_page_break()

# ── 6. FILE STRUCTURE ────────────────────────────────────────────────────────
heading1("6. Project Structure")

code("""nexus/
├── agents/
│   ├── orchestrator.py       # LangGraph graph definitions
│   ├── intake_agent.py       # PDF parsing + Claude extraction
│   ├── care_plan_agent.py    # OpenFDA + Claude care plan + Nebius
│   ├── monitoring_agent.py   # Daily check-in classification
│   ├── escalation_agent.py   # RED flag tier-based response
│   └── admin_agent.py        # Reminders + summaries
├── tools/
│   ├── pdf_parser.py         # PyMuPDF text extraction
│   ├── pinecone_store.py     # Vector store read/write
│   ├── openfda.py            # Medication interaction checks
│   ├── elevenlabs_stt.py     # Voice transcription
│   └── notification_tool.py  # Email + SMS
├── state/
│   └── recovery_state.py     # RecoveryState TypedDict schema
├── ui/
│   └── streamlit_app.py      # 7-page Streamlit UI
├── data/
│   └── test_discharge.pdf    # Synthetic CHF patient PDF
├── n8n/
│   └── daily_checkin_workflow.json
├── .streamlit/
│   └── config.toml           # Soft purple theme
├── .env.example
└── requirements.txt""")

doc.add_page_break()

# ── 7. UI PAGES ──────────────────────────────────────────────────────────────
heading1("7. Streamlit UI — 7 Pages")

table_2col([
    ("Onboarding", "Patient name, discharge date, emergency contact consent, PDF upload. Triggers the intake graph."),
    ("Care Plan", "4 tabs: First 7 Days checklist, Medications, Warning Signs, Appointments. Shows medication flags."),
    ("Daily Check-in", "Voice (ElevenLabs) or typed form. Individual medication checkboxes. Weight input step=1 lb."),
    ("Approvals", "Human review queue for medication flags and escalation message drafts before sending."),
    ("Dashboard", "Green/yellow/red day counts, daily vitals log, medication adherence progress bars."),
    ("Provider Summary", "Auto-generated brief with vitals trend, adherence, flagged anomalies. Exportable as .txt."),
    ("Hospital History", "Auto-populated from discharge PDF. Manually extensible for past hospitalizations."),
], header=("Page", "Description"))

doc.add_page_break()

# ── 8. KNOWN ISSUES & FIXES ──────────────────────────────────────────────────
heading1("8. Build Issues Encountered & Fixes")

heading2("Claude Returns Markdown-Wrapped JSON")
body(
    "All three agents (Intake, Care Plan, Monitoring) use Claude to return structured JSON. "
    "Despite prompts saying 'Return ONLY valid JSON', Claude sometimes wraps the response "
    "in ```json ... ``` code fences. Fix: strip code fences before calling json.loads() "
    "in all three agents."
)

heading2("Python Version Incompatibility")
body(
    "LangGraph requires Python 3.9+ for the dict[str, Any] generic type syntax. "
    "The initial Anaconda environment was Python 3.8. Fix: installed Python 3.11.9 "
    "and recreated the venv."
)

heading2("Pinecone Package Rename")
body(
    "The Pinecone Python package was renamed from pinecone-client to pinecone. "
    "Fix: pip uninstall pinecone-client && pip install pinecone."
)

heading2("Pinecone Index Dimensions")
body(
    "The build guide specified 1536 dimensions but Pinecone only offered 1024 or 2048. "
    "Fix: use 1024 dimensions. Updated embed_text() in pinecone_store.py to generate "
    "1024-dim vectors instead of 1536."
)

heading2("Helper Functions Inside if/elif Chain")
body(
    "_run_monitoring() and generate_provider_summary_text() were defined in the middle "
    "of Streamlit's if/elif page routing chain, causing SyntaxError. "
    "Fix: moved both functions to before the if/elif chain."
)

heading2("Missing Imports in intake_agent.py")
body(
    "datetime and uuid were used in run_intake_agent() but not imported at the top of "
    "the file. Fix: added 'import uuid' and 'from datetime import datetime'."
)

heading2("Nebius Model Name")
body(
    "The model ID meta-llama/Meta-Llama-3.1-70B-Instruct returned 404 on Nebius. "
    "Fix: updated to meta-llama/Meta-Llama-3.1-70B-Instruct-fast."
)

heading2("doc.close() Before len(doc) in pdf_parser.py")
body(
    "PyMuPDF raises 'document closed' if you call len(doc) after doc.close(). "
    "Fix: save page_count = len(doc) before calling doc.close()."
)

doc.add_page_break()

# ── 9. RUBRIC COMPLIANCE ─────────────────────────────────────────────────────
heading1("9. Week 3 Rubric Compliance")

table_2col([
    ("Multi-agent system", "5 specialized agents + orchestrator via LangGraph"),
    ("Tool use", "6 tools: PDF parser, Pinecone store, OpenFDA, ElevenLabs STT, notification, calendar (stub)"),
    ("Persistent state", "Pinecone vector store persists check-ins and discharge data across 30 days"),
    ("Error recovery", "JSON parse fallbacks, PDF failure routing, STT fallback to typed form, Pinecone local fallback"),
    ("Human-in-the-loop", "3 handoff types: medication flags, escalation messages, emergency contact consent"),
    ("Nebius Token Factory", "Parallel care plan generation via meta-llama/Meta-Llama-3.1-70B-Instruct-fast"),
    ("LangGraph", "StateGraph with conditional edges for both intake and monitoring graphs"),
    ("Voice input", "ElevenLabs Scribe v1 STT with Claude transcript parsing"),
], header=("Requirement", "Implementation"))

doc.add_page_break()

# ── 10. ENVIRONMENT SETUP ────────────────────────────────────────────────────
heading1("10. Environment Setup")

heading2("Prerequisites")
bullet("Python 3.11.9")
bullet("Git Bash (Windows)")
bullet("Pinecone account — index named 'nexus', 1024 dimensions, cosine metric")
bullet("Anthropic API key")
bullet("Nebius AI Studio API key")
bullet("ElevenLabs API key (STT access)")
bullet("Twilio account (SMS notifications)")
bullet("Gmail App Password (email notifications)")

heading2("Run Locally")
code("git clone https://github.com/raghavanlakshmis-hash/nexus.git")
code("cd nexus")
code("python -m venv venv")
code("source venv/Scripts/activate  # Git Bash on Windows")
code("pip install -r requirements.txt")
code("cp .env.example .env  # fill in your API keys")
code("streamlit run ui/streamlit_app.py")

heading2("Demo")
body(
    "Upload data/test_discharge.pdf on the onboarding page. This is a synthetic CHF "
    "patient discharge summary. No real patient data is used anywhere in the project."
)

# ── SAVE ─────────────────────────────────────────────────────────────────────
doc.save("Nexus_Project_Documentation.docx")
print("Nexus_Project_Documentation.docx created successfully")
